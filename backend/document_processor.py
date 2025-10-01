import os
import logging
from typing import Dict, List, Optional
from docx import Document
import PyPDF2
import pdfplumber
import json
from embedding_service import EmbeddingService
from vector_db_service import VectorDBService
from exact_match_service import ExactMatchService
from section_chunker import SectionChunker

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handle document text extraction and embedding-based processing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
        self.embedding_service = EmbeddingService()
        self.vector_db_service = VectorDBService()
        self.exact_match_service = ExactMatchService()
        self.section_chunker = SectionChunker()
    
    def extract_text(self, filepath: str) -> Optional[Dict]:
        """
        Extract text content from uploaded document
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return None
            
            file_extension = os.path.splitext(filepath)[1].lower()
            filename = os.path.basename(filepath)
            
            if file_extension == '.pdf':
                text_content = self._extract_pdf_text(filepath)
            elif file_extension == '.docx':
                text_content = self._extract_docx_text(filepath)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            if not text_content or not text_content.strip():
                logger.warning(f"No text extracted from {filename}")
                return {
                    "filename": filename,
                    "text_content": "",
                    "word_count": 0,
                    "char_count": 0,
                    "file_type": file_extension[1:],
                    "chunks": [],
                    "chunk_count": 0,
                    "error": "No text could be extracted from this document. The PDF may be image-based or corrupted."
                }
            
            # Generate chunks using embeddings
            chunks = self.embedding_service.create_chunks_with_embeddings(text_content, filename)
            
            return {
                "filename": filename,
                "filepath": filepath,
                "text_content": text_content,
                "word_count": len(text_content.split()),
                "char_count": len(text_content),
                "file_type": file_extension[1:],  # Remove the dot
                "chunks": chunks,
                "chunk_count": len(chunks)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return None
    
    def _extract_pdf_text(self, filepath: str) -> str:
        """Extract text from PDF file using multiple methods"""
        text = ""
        
        # Method 1: Try PyPDF2 first (faster for simple PDFs)
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted text using PyPDF2 from {filepath}")
                return text.strip()
                
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {filepath}: {str(e)}")
        
        # Method 2: Try pdfplumber (better for complex PDFs)
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted text using pdfplumber from {filepath}")
                return text.strip()
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {filepath}: {str(e)}")
        
        # Method 3: Try pdfplumber with layout preservation
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    # Try to extract text with layout preservation
                    page_text = page.extract_text(layout=True)
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted text using pdfplumber layout from {filepath}")
                return text.strip()
                
        except Exception as e:
            logger.warning(f"pdfplumber layout extraction failed for {filepath}: {str(e)}")
        
        logger.error(f"All PDF extraction methods failed for {filepath}")
        return ""
    
    def _extract_docx_text(self, filepath: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(filepath)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading DOCX {filepath}: {str(e)}")
            return ""
    
    def process_multiple_files(self, filepaths: List[str]) -> List[Dict]:
        """
        Process multiple documents and extract text
        
        Args:
            filepaths: List of file paths to process
            
        Returns:
            List of dictionaries containing extracted text and metadata
        """
        processed_docs = []
        
        for filepath in filepaths:
            doc_data = self.extract_text(filepath)
            if doc_data:
                processed_docs.append(doc_data)
                logger.info(f"Successfully processed: {doc_data['filename']}")
            else:
                logger.warning(f"Failed to process: {filepath}")
        
        return processed_docs
    
    def process_and_store_document(self, filepath: str) -> Dict:
        """
        Process document and store chunks in vector database
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Dictionary with processing results and vector DB storage info
        """
        try:
            # Extract text and create chunks using existing logic
            doc_data = self.extract_text(filepath)
            if not doc_data:
                raise ValueError(f"Failed to extract text from {filepath}")
            
            # Create chunks with embeddings (using section-based embedding service)
            chunks_with_embeddings = self.embedding_service.create_chunks_with_embeddings(
                doc_data['text_content'], 
                doc_data['filename']
            )
            
            if not chunks_with_embeddings:
                raise ValueError(f"Failed to create chunks for {doc_data['filename']}")
            
            # Prepare chunks for vector database
            vector_chunks = []
            for chunk_data in chunks_with_embeddings:
                vector_chunk = {
                    'text': chunk_data['text'],
                    'metadata': {
                        'start_char': chunk_data.get('start_char', 0),
                        'end_char': chunk_data.get('end_char', 0),
                        'char_count': chunk_data.get('char_count', len(chunk_data['text'])),
                        'word_count': len(chunk_data['text'].split()) if chunk_data['text'] else 0,
                        'chunk_type': 'semantic',
                        'file_extension': os.path.splitext(filepath)[1],
                        'processing_timestamp': doc_data.get('processing_timestamp')
                    }
                }
                vector_chunks.append(vector_chunk)
            
            # Store chunks in vector database
            chunk_ids = self.vector_db_service.add_document_chunks(
                document_name=doc_data['filename'],
                chunks=vector_chunks
            )
            
            # Add chunks to exact matching service (use section-based chunks with hashes)
            section_chunks = self.section_chunker.create_section_chunks(doc_data['text_content'], doc_data['filename'])
            self.exact_match_service.add_document_chunks(
                document_name=doc_data['filename'],
                chunks=section_chunks
            )
            
            # Update doc_data with vector database info
            doc_data.update({
                'vector_db_stored': True,
                'vector_chunk_ids': chunk_ids,
                'vector_chunk_count': len(chunk_ids),
                'chunks_with_embeddings': chunks_with_embeddings,  # Keep for compatibility
                'exact_match_enabled': True
            })
            
            logger.info(f"Successfully processed and stored '{doc_data['filename']}' "
                       f"with {len(chunk_ids)} chunks in vector database")
            
            return doc_data
            
        except Exception as e:
            logger.error(f"Failed to process and store document {filepath}: {str(e)}")
            raise
    
    def get_document_from_vector_db(self, document_name: str) -> Optional[Dict]:
        """
        Retrieve document chunks from vector database
        
        Args:
            document_name: Name of the document
            
        Returns:
            Dictionary with document data and chunks
        """
        try:
            chunks = self.vector_db_service.get_document_chunks(document_name)
            if not chunks:
                return None
            
            # Reconstruct document data
            doc_data = {
                'filename': document_name,
                'vector_db_stored': True,
                'vector_chunk_count': len(chunks),
                'chunks': chunks,
                'text_content': ' '.join([chunk['text'] for chunk in chunks])
            }
            
            # Add metadata from first chunk if available
            if chunks and chunks[0]['metadata']:
                first_metadata = chunks[0]['metadata']
                doc_data.update({
                    'file_extension': first_metadata.get('file_extension'),
                    'processing_timestamp': first_metadata.get('processing_timestamp')
                })
            
            logger.info(f"Retrieved '{document_name}' with {len(chunks)} chunks from vector database")
            return doc_data
            
        except Exception as e:
            logger.error(f"Failed to retrieve document from vector database: {str(e)}")
            return None
    
    def get_vector_db_stats(self) -> Dict:
        """Get statistics about the vector database"""
        try:
            return self.vector_db_service.get_collection_stats()
        except Exception as e:
            logger.error(f"Failed to get vector database stats: {str(e)}")
            raise
    
    def list_documents_in_vector_db(self) -> List[str]:
        """Get list of all documents in vector database"""
        try:
            return self.vector_db_service.list_documents()
        except Exception as e:
            logger.error(f"Failed to list documents: {str(e)}")
            raise
    
    def delete_document_from_vector_db(self, document_name: str) -> int:
        """
        Delete a document and all its chunks from vector database
        
        Args:
            document_name: Name of the document to delete
            
        Returns:
            Number of chunks deleted
        """
        try:
            # Remove from vector database
            deleted_count = self.vector_db_service.delete_document(document_name)
            
            # Remove from exact matching service
            self.exact_match_service.remove_document(document_name)
            
            logger.info(f"Deleted {deleted_count} chunks for document '{document_name}'")
            return deleted_count
        except Exception as e:
            logger.error(f"Failed to delete document from vector database: {str(e)}")
            raise
    
    def find_exact_matches(self, document_name: str) -> Dict:
        """Find exact matches for a specific document"""
        try:
            return self.exact_match_service.find_exact_matches(document_name)
        except Exception as e:
            logger.error(f"Failed to find exact matches: {str(e)}")
            raise
    
    def find_duplicate_sections(self) -> Dict:
        """Find all duplicate sections across documents"""
        try:
            return self.exact_match_service.find_duplicate_sections()
        except Exception as e:
            logger.error(f"Failed to find duplicate sections: {str(e)}")
            raise
    
    def compare_documents_exact(self, doc1_name: str, doc2_name: str) -> Dict:
        """Compare two documents for exact matches"""
        try:
            return self.exact_match_service.compare_documents_exact(doc1_name, doc2_name)
        except Exception as e:
            logger.error(f"Failed to compare documents: {str(e)}")
            raise
    
    def get_exact_match_stats(self) -> Dict:
        """Get statistics about exact matching"""
        try:
            return self.exact_match_service.get_document_stats()
        except Exception as e:
            logger.error(f"Failed to get exact match stats: {str(e)}")
            raise
    
    def compare_documents_sentence_level(self, doc1_name: str, doc2_name: str) -> Dict:
        """Compare two documents for exact sentence matches"""
        try:
            return self.exact_match_service.compare_documents_sentence_level(doc1_name, doc2_name)
        except Exception as e:
            logger.error(f"Failed to compare documents at sentence level: {str(e)}")
            raise
